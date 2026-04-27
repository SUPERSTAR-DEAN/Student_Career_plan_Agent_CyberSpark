# -*- coding: utf-8 -*-
"""职业生涯发展报告生成器（摘要、路径、行动计划、导出）"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from models.llm_wrapper import LLMWrapper


def _format_action_plan_lines(items: Any) -> str:
    """行动计划列表转为统一正文条目（无前导 # / -）。"""
    if not items:
        return "  · （暂无）"
    lines = []
    for t in items:
        if not isinstance(t, dict):
            continue
        task = (t.get("task") or "").strip()
        dl = (t.get("deadline") or "").strip()
        lines.append(f"  · {task}（{dl}）" if dl else f"  · {task}")
    return "\n".join(lines) if lines else "  · （暂无）"


def _format_action_plan_lines_with_resources(items: Any) -> str:
    """行动计划转正文条目，并附带资源链接文本。"""
    if not items:
        return "  · （暂无）"
    lines = []
    for t in items:
        if not isinstance(t, dict):
            continue
        task = (t.get("task") or "").strip()
        dl = (t.get("deadline") or "").strip()
        head = f"  · {task}（{dl}）" if dl else f"  · {task}"
        lines.append(head)
        for r in (t.get("resources") or [])[:4]:
            if not isinstance(r, dict):
                lines.append(f"    - 资源：{str(r)}")
                continue
            name = r.get("name", "资源")
            url = r.get("url", "")
            reason = r.get("reason", "")
            lines.append(f"    - 资源：{name} {url}".strip())
            if reason:
                lines.append(f"      用法：{reason}")
    return "\n".join(lines) if lines else "  · （暂无）"


def _mid_term_paths_display(mid: Any) -> str:
    """将中期路径列表格式化为可读中文（避免 full_text 里出现 Python 列表或裸 job_id）。"""
    if mid is None:
        return "（暂无）"
    if isinstance(mid, str):
        return mid.strip() or "（暂无）"
    if isinstance(mid, list):
        if not mid:
            return "（暂无）"
        return "、".join(str(x).strip() for x in mid if str(x).strip())
    return str(mid)


def _job_id_to_name(job_profiles: dict, job_id: str) -> str:
    prof = job_profiles.get(job_id) if job_profiles else None
    if prof is not None:
        return getattr(prof, "job_name", str(job_id))
    return str(job_id)


def _infer_plan_track(target_job: str, gap_analysis: dict, top_match: Optional[dict] = None) -> str:
    text = f"{target_job}\n{(top_match or {}).get('job_name','')}\n" + " ".join(
        str(x) for x in (gap_analysis.get("missing_skills") or [])
    )
    t = text.lower()
    if any(k in t for k in ("科研", "研究员", "算法研究", "论文", "实验")):
        return "research"
    if any(k in t for k in ("前端", "react", "vue", "javascript", "typescript", "h5", "web")):
        return "frontend"
    if any(k in t for k in ("后端", "服务端", "spring", "java", "fastapi", "django", "go")):
        return "backend"
    if any(k in t for k in ("数据库", "dba", "oracle", "mysql", "postgres", "sql")):
        return "database"
    if any(k in t for k in ("数据开发", "数据工程", "etl", "hadoop", "spark", "flink", "kafka", "数仓")):
        return "data"
    return "general"


def _resource_pool(track: str) -> dict[str, list[dict]]:
    """离线可用的高质量资源池（论坛 + 课程 + B站），用于生成可执行计划。"""
    common_forum = [
        {"name": "掘金", "type": "论坛", "url": "https://juejin.cn/", "reason": "看工程实战与岗位技能路线"},
        {"name": "CSDN", "type": "论坛", "url": "https://www.csdn.net/", "reason": "问题检索速度快，适合排障"},
        {"name": "V2EX", "type": "社区", "url": "https://www.v2ex.com/", "reason": "关注行业趋势与真实工作话题"},
    ]
    common_video = [
        {"name": "B站·黑马程序员", "type": "B站课程", "url": "https://space.bilibili.com/37974444", "reason": "体系化课程，适合打基础"},
        {"name": "B站·尚硅谷", "type": "B站课程", "url": "https://space.bilibili.com/302417610", "reason": "Java/数据库/大数据课程完整"},
        {"name": "B站·CodeSheep", "type": "B站UP主", "url": "https://space.bilibili.com/384068749", "reason": "职业成长与技术认知"},
    ]
    by_track = {
        "frontend": [
            {"name": "MDN Web Docs", "type": "文档", "url": "https://developer.mozilla.org/zh-CN/", "reason": "前端标准规范权威资料"},
            {"name": "B站·前端开发教程检索", "type": "B站搜索", "url": "https://search.bilibili.com/all?keyword=React%20Vue%20TypeScript%20%E9%A1%B9%E7%9B%AE%E5%AE%9E%E6%88%98", "reason": "按 React/Vue/TS 实战搜课程"},
        ],
        "backend": [
            {"name": "Spring 官方文档", "type": "文档", "url": "https://spring.io/projects", "reason": "后端框架能力提升主线"},
            {"name": "B站·后端项目实战检索", "type": "B站搜索", "url": "https://search.bilibili.com/all?keyword=Java%20SpringBoot%20%E5%90%8E%E7%AB%AF%E9%A1%B9%E7%9B%AE%E5%AE%9E%E6%88%98", "reason": "项目链路完整（接口/鉴权/部署）"},
        ],
        "database": [
            {"name": "MySQL 官方文档", "type": "文档", "url": "https://dev.mysql.com/doc/", "reason": "索引/事务/优化必须掌握"},
            {"name": "B站·数据库优化检索", "type": "B站搜索", "url": "https://search.bilibili.com/all?keyword=MySQL%20%E7%B4%A2%E5%BC%95%20SQL%E4%BC%98%E5%8C%96", "reason": "集中补齐 SQL 与调优能力"},
        ],
        "data": [
            {"name": "Apache Spark 文档", "type": "文档", "url": "https://spark.apache.org/docs/latest/", "reason": "数据开发与数仓关键组件"},
            {"name": "B站·数仓与ETL检索", "type": "B站搜索", "url": "https://search.bilibili.com/all?keyword=%E6%95%B0%E4%BB%93%20ETL%20Spark%20Flink", "reason": "从入门到工程化链路学习"},
        ],
        "research": [
            {"name": "Google Scholar", "type": "论文检索", "url": "https://scholar.google.com/", "reason": "快速定位领域经典论文"},
            {"name": "arXiv", "type": "论文平台", "url": "https://arxiv.org/", "reason": "跟踪前沿研究方向"},
            {"name": "B站·论文阅读与科研方法检索", "type": "B站搜索", "url": "https://search.bilibili.com/all?keyword=%E8%AE%BA%E6%96%87%E9%98%85%E8%AF%BB%20%E7%A7%91%E7%A0%94%E6%96%B9%E6%B3%95", "reason": "提升论文阅读与实验复现效率"},
        ],
        "general": [
            {"name": "LeetCode", "type": "刷题", "url": "https://leetcode.cn/", "reason": "补齐算法与问题解决能力"},
            {"name": "B站·计算机基础检索", "type": "B站搜索", "url": "https://search.bilibili.com/all?keyword=%E8%AE%A1%E7%AE%97%E6%9C%BA%E5%9F%BA%E7%A1%80%20%E6%95%B0%E6%8D%AE%E7%BB%93%E6%9E%84%20%E7%AE%97%E6%B3%95", "reason": "系统补基础，提升迁移能力"},
        ],
    }
    extra = by_track.get(track, by_track["general"])
    return {"forums": common_forum, "videos": common_video, "extra": extra}


class CareerReportGenerator:
    """职业生涯发展报告生成器"""

    def __init__(self, llm_wrapper: Optional[LLMWrapper] = None):
        self.llm = llm_wrapper

    def generate_executive_summary(self, match_result: dict) -> str:
        """生成报告摘要：优势、差距、核心建议"""
        overall = match_result.get("overall_score", 0)
        gap = match_result.get("gap_analysis", {})
        advantage = gap.get("advantage_skills", [])
        missing = gap.get("missing_skills", [])
        to_improve = gap.get("to_improve", {})
        # 各维度得分由前端以柱状图展示，摘要中不再输出英文 key 的字典形式
        parts = [f"综合人岗匹配度：{overall}分。"]
        if advantage:
            parts.append(f"当前优势技能：{', '.join(advantage[:10])}。")
        if missing:
            parts.append(f"建议补足技能：{', '.join(missing[:10])}。")
        if to_improve:
            parts.append(f"建议重点提升：{', '.join(list(to_improve.keys())[:5])}。")
        summary = " ".join(parts)
        if self.llm:
            try:
                llm_text = self.llm.generate_career_advice({"summary": summary, "match_result": match_result})
                # 若模型返回的是 JSON/结构化内容，仍使用已生成的可读摘要
                if isinstance(llm_text, str) and llm_text.strip() and not llm_text.strip().startswith(("{", "[")):
                    summary = llm_text
            except Exception:
                pass
        return summary

    def generate_career_path_section(
        self,
        target_job: str,
        graph_builder: Any,
    ) -> dict:
        """生成职业路径规划章节：短期、中期、长期、路径图数据"""
        vertical_paths = []
        lateral_paths = []
        path_sequence = []
        if graph_builder:
            jid = getattr(graph_builder, "_resolve_job_id", lambda x: x)(target_job) or target_job
            vertical_paths = getattr(graph_builder, "get_vertical_paths", lambda x: [])(jid)
            lateral_paths = getattr(graph_builder, "get_lateral_paths", lambda x: [])(jid)
            path_sequence = getattr(graph_builder, "get_career_path", lambda s, t: [])(target_job, None)
        job_profiles = getattr(graph_builder, "job_profiles", {}) or {}
        short_term = target_job
        mid_ids: list[str] = []
        for p in (vertical_paths or [])[:1]:
            if len(p) >= 2:
                mid_ids.append(p[1])
        for p in (lateral_paths or [])[:2]:
            if len(p) >= 2:
                mid_ids.append(p[1])
        mid_ids = list(dict.fromkeys(mid_ids))
        mid_path = [_job_id_to_name(job_profiles, jid) for jid in mid_ids]
        long_vision = "成为该领域资深专家或管理岗位" if mid_path else "持续晋升至更高层级"
        return {
            "short_term": short_term,
            "mid_term_paths": mid_path,
            "long_term_vision": long_vision,
            "vertical_paths": vertical_paths,
            "lateral_paths": lateral_paths,
            "path_sequence": path_sequence,
            "graph_data": getattr(graph_builder, "visualize_graph", lambda: {"nodes": [], "edges": []})(),
        }

    def generate_action_plan(
        self,
        gap_analysis: dict,
        target_job: str = "",
        top_match: Optional[dict] = None,
        timeline: str = "6_months",
    ) -> dict:
        """生成分阶段行动计划：短期/中期/长期 + 具体资源推荐 + 口语化段落说明。"""
        _ = timeline
        short_term = []
        mid_term = []
        long_term = []
        missing = gap_analysis.get("missing_skills", [])[:5]
        to_improve = gap_analysis.get("to_improve", {})
        missing_certs = gap_analysis.get("missing_certificates", [])[:3]
        track = _infer_plan_track(target_job, gap_analysis, top_match=top_match)
        res = _resource_pool(track)

        for i, skill in enumerate(missing):
            short_term.append({
                "task": f"学习或巩固：{skill}（做 1 个可展示的小项目并写复盘）",
                "deadline": f"{(i + 1)}个月",
                "resources": [res["extra"][0], res["videos"][0], res["forums"][0]],
            })
        for skill, level in list(to_improve.items())[:3]:
            short_term.append({
                "task": f"提升 {skill} 至熟练（从当前 {level}/5 提升到 4/5）",
                "deadline": "2个月",
                "resources": [res["videos"][1], res["forums"][1]],
            })
        for c in missing_certs:
            mid_term.append({
                "task": f"考取/准备：{c}",
                "deadline": "6个月内",
                "resources": [res["forums"][1], {"name": "中国教育考试网", "type": "官网", "url": "https://www.neea.edu.cn/", "reason": "报名信息与考试说明"}],
            })
        mid_term.append({
            "task": "争取对口实习或真实业务项目（至少 1 段）",
            "deadline": "6个月内",
            "resources": [
                {"name": "牛客校招", "type": "平台", "url": "https://www.nowcoder.com/", "reason": "校招与实习信息集中"},
                {"name": "BOSS直聘校招", "type": "平台", "url": "https://www.zhipin.com/", "reason": "岗位覆盖广，投递反馈快"},
                res["forums"][2],
            ],
        })
        long_term.extend(
            [
                {
                    "task": "形成“方向标签 + 代表项目 + 可量化结果”的个人品牌",
                    "deadline": "12个月",
                    "resources": [res["extra"][0], res["forums"][0]],
                },
                {
                    "task": f"围绕目标岗位「{target_job or '目标方向'}」做一次系统性能力盘点与路径升级",
                    "deadline": "12-18个月",
                    "resources": [res["extra"][-1], res["videos"][-1]],
                },
            ]
        )

        short_para = (
            "先别急着把战线拉太长，前 1-2 个月只做一件事：把岗位里最缺的 2-3 个核心能力补齐，"
            "每个能力都要有作品或项目证明。建议你按“学一块 -> 做一块 -> 讲一块”的节奏推进，"
            "这样面试时不只是会背概念，而是能说清楚你是怎么解决问题的。"
        )
        mid_para = (
            "中期目标是把能力从“会用”升级到“可交付”。你需要至少一段对口实习或真实项目经历，"
            "并且能量化结果，比如性能提升、缺陷率下降、上线周期缩短。证书不是必须，但对部分岗位是加分项，"
            "建议与投递节奏并行推进。"
        )
        long_para = (
            "长期看，核心不是多会几门技术，而是形成稳定的职业定位。你可以把自己定义成“某方向 + 某场景 + 某结果”的人才，"
            "例如“偏后端的性能优化型工程师”或“偏科研落地的算法工程师”。定位清楚后，成长速度会明显快于盲目跟风。"
        )

        if self.llm:
            try:
                # 用 LLM 做轻量润色：口语化 + 专业化并存
                ctx = {
                    "target_job": target_job,
                    "track": track,
                    "missing_skills": missing,
                    "to_improve": to_improve,
                    "base_text": f"{short_para}\n{mid_para}\n{long_para}",
                }
                p = self.llm.generate_career_advice(ctx)
                if isinstance(p, str) and p.strip() and not p.strip().startswith(("{", "[")):
                    short_para = p.strip()
            except Exception:
                pass

        return {
            "short_term": short_term,
            "mid_term": mid_term,
            "long_term": long_term,
            "narrative": {
                "short_term_paragraph": short_para,
                "mid_term_paragraph": mid_para,
                "long_term_paragraph": long_para,
                "track": track,
            },
            "evaluation_cycle": "每月自评一次技能与进展",
            "evaluation_metrics": ["技能掌握进度", "实习/项目投递与反馈", "证书备考进度"],
        }

    def compile_full_report(
        self,
        student_profile: Any,
        match_results: list[dict],
        career_path: dict,
        action_plan: dict,
        executive_summary: Optional[str] = None,
    ) -> dict:
        """整合生成完整结构化报告"""
        top_match = match_results[0] if match_results else {}
        if executive_summary is None:
            executive_summary = self.generate_executive_summary(top_match)
        target_job = top_match.get("job_name", "")
        report = {
            "title": "大学生职业生涯发展报告",
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "student_id": getattr(student_profile, "student_id", ""),
            "completeness_score": getattr(student_profile, "completeness_score", 0),
            "competitiveness_score": getattr(student_profile, "competitiveness_score", 0),
            "executive_summary": executive_summary,
            "match_results": match_results[:10],
            "top_job": target_job,
            "career_path": career_path,
            "action_plan": action_plan,
            "full_text": "",
        }
        short_lines = _format_action_plan_lines_with_resources(action_plan.get("short_term", []))
        mid_lines = _format_action_plan_lines_with_resources(action_plan.get("mid_term", []))
        long_lines = _format_action_plan_lines_with_resources(action_plan.get("long_term", []))
        nv = action_plan.get("narrative", {}) if isinstance(action_plan, dict) else {}
        full_parts = [
            f"{report['title']}\n",
            f"生成时间：{report['generated_at']}\n",
            f"\n一、摘要\n{(executive_summary or '').strip()}\n",
            f"\n二、人岗匹配概览\n推荐岗位：{target_job or '（暂无）'}，匹配度：{top_match.get('overall_score', 0)}分。\n",
            "\n三、职业路径规划\n"
            f"短期目标：{career_path.get('short_term', '')}\n"
            f"中期路径：{_mid_term_paths_display(career_path.get('mid_term_paths'))}\n"
            f"长期愿景：{career_path.get('long_term_vision', '')}\n",
            "\n四、行动计划\n"
            f"\n【短期策略说明】\n{nv.get('short_term_paragraph', '')}\n"
            "\n短期计划\n"
            f"{short_lines}\n"
            f"\n【中期策略说明】\n{nv.get('mid_term_paragraph', '')}\n"
            "\n中期计划\n"
            f"{mid_lines}\n"
            f"\n【长期策略说明】\n{nv.get('long_term_paragraph', '')}\n"
            "\n长期计划\n"
            f"{long_lines}\n\n"
            f"评估周期：{action_plan.get('evaluation_cycle', '')}\n"
            f"评估指标：{'、'.join(action_plan.get('evaluation_metrics', []) or [])}",
        ]
        report["full_text"] = "\n".join(full_parts)
        return report

    def export_to_pdf(self, report_data: dict, output_path: str) -> bool:
        """导出报告为 PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(str(path), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            try:
                pdfmetrics.registerFont(TTFont("SimSun", "simsun.ttc"))
                style = ParagraphStyle("Custom", fontName="SimSun", fontSize=10)
            except Exception:
                style = styles["Normal"]
            text = report_data.get("full_text", "")
            for line in text.split("\n"):
                if line.strip():
                    story.append(Paragraph(line.replace("<", "&lt;").replace(">", "&gt;"), style))
                    story.append(Spacer(1, 4 * mm))
            doc.build(story)
            return True
        except Exception:
            return False

    def export_to_word(self, report_data: dict, output_path: str) -> bool:
        """导出报告为 Word"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_heading(report_data.get("title", "职业生涯发展报告"), 0)
            doc.add_paragraph(f"生成时间：{report_data.get('generated_at', '')}")
            doc.add_paragraph(f"完整度评分：{report_data.get('completeness_score', 0)}；竞争力评分：{report_data.get('competitiveness_score', 0)}")
            doc.add_heading("一、摘要", level=1)
            doc.add_paragraph(report_data.get("executive_summary", ""))
            doc.add_heading("二、人岗匹配", level=1)
            for r in report_data.get("match_results", [])[:5]:
                doc.add_paragraph(f"{r.get('job_name', '')} — 匹配度 {r.get('overall_score', 0)} 分")
            doc.add_heading("三、职业路径", level=1)
            cp = report_data.get("career_path", {})
            doc.add_paragraph(f"短期：{cp.get('short_term', '')}")
            doc.add_paragraph(f"中期：{_mid_term_paths_display(cp.get('mid_term_paths'))}")
            doc.add_paragraph(f"长期：{cp.get('long_term_vision', '')}")
            doc.add_heading("四、行动计划", level=1)
            ap = report_data.get("action_plan", {})
            for t in ap.get("short_term", []):
                doc.add_paragraph(f"• {t.get('task', '')}（{t.get('deadline', '')}）", style="List Bullet")
            for t in ap.get("mid_term", []):
                doc.add_paragraph(f"• {t.get('task', '')}（{t.get('deadline', '')}）", style="List Bullet")
            doc.add_paragraph(f"评估周期：{ap.get('evaluation_cycle', '')}")
            doc.save(str(path))
            return True
        except Exception:
            return False
