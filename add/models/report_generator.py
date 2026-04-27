# -*- coding: utf-8 -*-
"""职业生涯发展报告生成器（摘要、路径、行动计划、导出）"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent))
from models.llm_wrapper import LLMWrapper

class CareerReportGenerator:
    """职业生涯发展报告生成器"""

    def __init__(self, llm_wrapper: Optional[LLMWrapper] = None):
        self.llm = llm_wrapper

    def generate_executive_summary(self, match_result: dict) -> str:
        """生成报告摘要：优势、差距、核心建议"""
        overall = match_result.get("overall_score", 0)
        dim_scores = match_result.get("dimension_scores", {})
        gap = match_result.get("gap_analysis", {})
        advantage = gap.get("advantage_skills", [])
        missing = gap.get("missing_skills", [])
        to_improve = gap.get("to_improve", {})
        parts = [
            f"综合人岗匹配度：{overall}分。",
            f"各维度得分：{dim_scores}。",
        ]
        if advantage:
            parts.append(f"当前优势技能：{', '.join(advantage[:10])}。")
        if missing:
            parts.append(f"建议补足技能：{', '.join(missing[:10])}。")
        if to_improve:
            parts.append(f"建议重点提升：{', '.join(list(to_improve.keys())[:5])}。")
        summary = " ".join(parts)
        if self.llm:
            try:
                llm_text = self.llm.generate_career_advice({"summary": summary, "match_result": match_result})
                # 若模型返回的是 JSON/结构化内容，仍使用已生成的可读摘要
                if isinstance(llm_text, str) and llm_text.strip() and not llm_text.strip().startswith(("{", "[")):
                    summary = llm_text
            except Exception:
                pass
        return summary

    def generate_career_path_section(
        self,
        target_job: str,
        graph_builder: Any,
    ) -> dict:
        """生成职业路径规划章节：短期、中期、长期、路径图数据"""
        vertical_paths = []
        lateral_paths = []
        path_sequence = []
        if graph_builder:
            jid = getattr(graph_builder, "_resolve_job_id", lambda x: x)(target_job) or target_job
            vertical_paths = getattr(graph_builder, "get_vertical_paths", lambda x: [])(jid)
            lateral_paths = getattr(graph_builder, "get_lateral_paths", lambda x: [])(jid)
            path_sequence = getattr(graph_builder, "get_career_path", lambda s, t: [])(target_job, None)
        job_profiles = getattr(graph_builder, "job_profiles", {})
        short_term = target_job
        mid_path = []
        for p in (vertical_paths or [])[:1]:
            if len(p) >= 2:
                mid_path.append(p[1])
        for p in (lateral_paths or [])[:2]:
            if len(p) >= 2:
                mid_path.append(p[1])
        mid_path = list(dict.fromkeys(mid_path))
        long_vision = "成为该领域资深专家或管理岗位" if mid_path else "持续晋升至更高层级"
        return {
            "short_term": short_term,
            "mid_term_paths": mid_path,
            "long_term_vision": long_vision,
            "vertical_paths": vertical_paths,
            "lateral_paths": lateral_paths,
            "path_sequence": path_sequence,
            "graph_data": getattr(graph_builder, "visualize_graph", lambda: {"nodes": [], "edges": []})(),
        }

    def generate_action_plan(
        self,
        gap_analysis: dict,
        timeline: str = "6_months",
    ) -> dict:
        """生成分阶段行动计划：短期、中期"""
        short_term = []
        mid_term = []
        missing = gap_analysis.get("missing_skills", [])[:5]
        to_improve = gap_analysis.get("to_improve", {})
        missing_certs = gap_analysis.get("missing_certificates", [])[:3]
        for i, skill in enumerate(missing):
            short_term.append({
                "task": f"学习或巩固：{skill}",
                "deadline": f"{(i + 1)}个月",
                "resources": ["在线课程", "项目实践"],
            })
        for skill, level in list(to_improve.items())[:3]:
            short_term.append({
                "task": f"提升{skill}至熟练",
                "deadline": "2个月",
                "resources": ["实战项目", "刷题/案例"],
            })
        for c in missing_certs:
            mid_term.append({
                "task": f"考取/准备：{c}",
                "deadline": "6个月内",
                "resources": ["官方教材", "真题"],
            })
        mid_term.append({
            "task": "争取对口实习或项目经历",
            "deadline": "6个月内",
            "resources": ["校招/实习平台", "导师项目"],
        })
        return {
            "short_term": short_term,
            "mid_term": mid_term,
            "evaluation_cycle": "每月自评一次技能与进展",
            "evaluation_metrics": ["技能掌握进度", "实习/项目投递与反馈", "证书备考进度"],
        }

    def compile_full_report(
        self,
        student_profile: Any,
        match_results: list[dict],
        career_path: dict,
        action_plan: dict,
        executive_summary: Optional[str] = None,
    ) -> dict:
        """整合生成完整结构化报告"""
        top_match = match_results[0] if match_results else {}
        if executive_summary is None:
            executive_summary = self.generate_executive_summary(top_match)
        target_job = top_match.get("job_name", "")
        report = {
            "title": "大学生职业生涯发展报告",
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "student_id": getattr(student_profile, "student_id", ""),
            "completeness_score": getattr(student_profile, "completeness_score", 0),
            "competitiveness_score": getattr(student_profile, "competitiveness_score", 0),
            "executive_summary": executive_summary,
            "match_results": match_results[:10],
            "top_job": target_job,
            "career_path": career_path,
            "action_plan": action_plan,
            "full_text": "",
        }
        full_parts = [
            f"# {report['title']}\n",
            f"生成时间：{report['generated_at']}\n",
            f"## 一、摘要\n{executive_summary}\n",
            f"## 二、人岗匹配概览\n推荐岗位：{target_job}，匹配度：{top_match.get('overall_score', 0)}分。\n",
            f"## 三、职业路径规划\n短期目标：{career_path.get('short_term', '')}\n中期路径：{career_path.get('mid_term_paths', [])}\n长期愿景：{career_path.get('long_term_vision', '')}\n",
            "## 四、行动计划\n",
            "### 短期\n" + "\n".join(f"- {t.get('task', '')}（{t.get('deadline', '')}）" for t in action_plan.get("short_term", [])),
            "\n### 中期\n" + "\n".join(f"- {t.get('task', '')}（{t.get('deadline', '')}）" for t in action_plan.get("mid_term", [])),
            f"\n评估周期：{action_plan.get('evaluation_cycle', '')}",
        ]
        report["full_text"] = "\n".join(full_parts)
        return report

    def export_to_pdf(self, report_data: dict, output_path: str) -> bool:
        """导出报告为 PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(str(path), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            try:
                pdfmetrics.registerFont(TTFont("SimSun", "simsun.ttc"))
                style = ParagraphStyle("Custom", fontName="SimSun", fontSize=10)
            except Exception:
                style = styles["Normal"]
            text = report_data.get("full_text", "")
            for line in text.split("\n"):
                if line.strip():
                    story.append(Paragraph(line.replace("<", "&lt;").replace(">", "&gt;"), style))
                    story.append(Spacer(1, 4 * mm))
            doc.build(story)
            return True
        except Exception:
            return False

    def export_to_word(self, report_data: dict, output_path: str) -> bool:
        """导出报告为 Word"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_heading(report_data.get("title", "职业生涯发展报告"), 0)
            doc.add_paragraph(f"生成时间：{report_data.get('generated_at', '')}")
            doc.add_paragraph(f"完整度评分：{report_data.get('completeness_score', 0)}；竞争力评分：{report_data.get('competitiveness_score', 0)}")
            doc.add_heading("一、摘要", level=1)
            doc.add_paragraph(report_data.get("executive_summary", ""))
            doc.add_heading("二、人岗匹配", level=1)
            for r in report_data.get("match_results", [])[:5]:
                doc.add_paragraph(f"{r.get('job_name', '')} — 匹配度 {r.get('overall_score', 0)} 分")
            doc.add_heading("三、职业路径", level=1)
            cp = report_data.get("career_path", {})
            doc.add_paragraph(f"短期：{cp.get('short_term', '')}")
            doc.add_paragraph(f"中期：{cp.get('mid_term_paths', [])}")
            doc.add_paragraph(f"长期：{cp.get('long_term_vision', '')}")
            doc.add_heading("四、行动计划", level=1)
            ap = report_data.get("action_plan", {})
            for t in ap.get("short_term", []):
                doc.add_paragraph(f"• {t.get('task', '')}（{t.get('deadline', '')}）", style="List Bullet")
            for t in ap.get("mid_term", []):
                doc.add_paragraph(f"• {t.get('task', '')}（{t.get('deadline', '')}）", style="List Bullet")
            doc.add_paragraph(f"评估周期：{ap.get('evaluation_cycle', '')}")
            doc.save(str(path))
            return True
        except Exception:
            return False
